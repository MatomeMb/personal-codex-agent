"""
Document Processor for Personal Codex Agent
Handles multiple file formats and extracts text for RAG implementation
"""

import os
import re
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path

# Document processing imports
try:
    import pypdf
    from docx import Document
    import markdown
except ImportError:
    print("Warning: Some document processing libraries not available")

from .exceptions import DocumentProcessingError

class DocumentProcessor:
    """Handles document loading, processing, and chunking for RAG implementation"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.logger = logging.getLogger(__name__)
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.supported_formats = ['.pdf', '.docx', '.md', '.txt']
        
        self.logger.info(f"DocumentProcessor initialized with chunk_size={chunk_size}, chunk_overlap={chunk_overlap}")
    
    def load_document(self, file_path: str) -> Dict[str, Any]:
        """Load a document and return its content and metadata"""
        file_path = Path(file_path)
        
        self.logger.info(f"Loading document: {file_path}")
        
        if not file_path.exists():
            self.logger.error(f"File not found: {file_path}")
            raise DocumentProcessingError(f"File not found: {file_path}")
        
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            self.logger.error(f"Unsupported file format: {file_extension}")
            raise DocumentProcessingError(f"Unsupported file format: {file_extension}. Supported formats: {', '.join(self.supported_formats)}")
        
        try:
            if file_extension == '.pdf':
                result = self._load_pdf(file_path)
            elif file_extension == '.docx':
                result = self._load_docx(file_path)
            elif file_extension == '.md':
                result = self._load_markdown(file_path)
            elif file_extension == '.txt':
                result = self._load_text(file_path)
            else:
                raise DocumentProcessingError(f"Unhandled file format: {file_extension}")
            
            self.logger.info(f"Successfully loaded document: {file_path}")
            return result
            
        except DocumentProcessingError:
            raise
        except Exception as e:
            self.logger.error(f"Unexpected error processing {file_path}: {e}")
            raise DocumentProcessingError(f"Error processing {file_path}: {str(e)}")
    
    def _load_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from PDF files"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                
                return {
                    'content': text,
                    'metadata': {
                        'source': str(file_path),
                        'type': 'pdf',
                        'pages': len(pdf_reader.pages),
                        'filename': file_path.name
                    }
                }
        except Exception as e:
            raise Exception(f"PDF processing error: {str(e)}")
    
    def _load_docx(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from Word documents"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return {
                'content': text,
                'metadata': {
                    'source': str(file_path),
                    'type': 'docx',
                    'paragraphs': len(doc.paragraphs),
                    'filename': file_path.name
                }
            }
        except Exception as e:
            raise Exception(f"DOCX processing error: {str(e)}")
    
    def _load_markdown(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from Markdown files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                # Convert markdown to plain text (remove formatting)
                text = markdown.markdown(content)
                # Remove HTML tags
                text = re.sub(r'<[^>]+>', '', text)
                
                return {
                    'content': text,
                    'metadata': {
                        'source': str(file_path),
                        'type': 'markdown',
                        'filename': file_path.name
                    }
                }
        except Exception as e:
            raise Exception(f"Markdown processing error: {str(e)}")
    
    def _load_text(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                
                return {
                    'content': content,
                    'metadata': {
                        'source': str(file_path),
                        'type': 'text',
                        'filename': file_path.name
                    }
                }
        except Exception as e:
            raise Exception(f"Text processing error: {str(e)}")
    
    def chunk_text(self, text: str, chunk_size: Optional[int] = None, 
                   chunk_overlap: Optional[int] = None) -> List[str]:
        """Split text into overlapping chunks for RAG processing"""
        if chunk_size is None:
            chunk_size = self.chunk_size
        if chunk_overlap is None:
            chunk_overlap = self.chunk_overlap
        
        # Clean and normalize text
        text = re.sub(r'\s+', ' ', text).strip()
        
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # If this isn't the last chunk, try to break at a sentence boundary
            if end < len(text):
                # Look for sentence endings within the last 100 characters
                search_start = max(start + chunk_size - 100, start)
                search_text = text[search_start:end]
                
                # Find the last sentence ending
                sentence_endings = ['.', '!', '?', '\n']
                last_ending = -1
                for ending in sentence_endings:
                    pos = search_text.rfind(ending)
                    if pos != -1:
                        last_ending = max(last_ending, pos)
                
                if last_ending != -1:
                    end = search_start + last_ending + 1
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move to next chunk with overlap
            start = end - chunk_overlap
            if start >= len(text):
                break
        
        return chunks
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """Process a document and return chunks with metadata"""
        document = self.load_document(file_path)
        chunks = self.chunk_text(document['content'])
        
        processed_chunks = []
        for i, chunk in enumerate(chunks):
            processed_chunks.append({
                'content': chunk,
                'chunk_id': i,
                'metadata': {
                    **document['metadata'],
                    'chunk_index': i,
                    'total_chunks': len(chunks)
                }
            })
        
        return {
            'document': document,
            'chunks': processed_chunks,
            'total_chunks': len(chunks)
        }
    
    def process_directory(self, directory_path: str) -> List[Dict[str, Any]]:
        """Process all supported documents in a directory"""
        directory = Path(directory_path)
        
        self.logger.info(f"Processing directory: {directory_path}")
        
        if not directory.exists():
            self.logger.error(f"Directory does not exist: {directory_path}")
            raise DocumentProcessingError(f"Directory does not exist: {directory_path}")
        
        if not directory.is_dir():
            self.logger.error(f"Path is not a directory: {directory_path}")
            raise DocumentProcessingError(f"Path is not a directory: {directory_path}")
        
        processed_documents = []
        errors = []
        
        for file_path in directory.iterdir():
            if file_path.is_file() and file_path.suffix.lower() in self.supported_formats:
                try:
                    self.logger.info(f"Processing file: {file_path}")
                    processed = self.process_document(str(file_path))
                    processed_documents.append(processed)
                    self.logger.info(f"Successfully processed: {file_path}")
                except DocumentProcessingError as e:
                    self.logger.error(f"Failed to process {file_path}: {e}")
                    errors.append(f"{file_path}: {e}")
                except Exception as e:
                    self.logger.error(f"Unexpected error processing {file_path}: {e}")
                    errors.append(f"{file_path}: {e}")
        
        if errors:
            self.logger.warning(f"Encountered {len(errors)} errors during processing")
            for error in errors:
                self.logger.warning(f"Processing error: {error}")
        
        self.logger.info(f"Processed {len(processed_documents)} documents successfully")
        return processed_documents
